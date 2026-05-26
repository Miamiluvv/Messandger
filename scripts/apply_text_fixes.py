# -*- coding: utf-8 -*-
"""Безопасные текстовые правки:
  1. Унифицируем тире: ' - ' и ' – ' → ' — ' (только в основном тексте,
     НЕ в подписях рисунков и НЕ в списке источников).
  2. ' (рисунок N)' → ', как показано на рисунке N' — единообразный стиль.
  3. ' (рисунок N) и ' → ', как показано на рисунке N, и '
"""
import os, re, shutil
from datetime import datetime
from docx import Document

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TARGET = os.path.join(ROOT, 'ДИПЛОМ_Махмудова_А_Д [HIGirw] [Mbuh0y].docx')


def make_backup():
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    backup = os.path.join(ROOT, f'_BACKUP_до_текстправок_{ts}.docx')
    shutil.copy2(TARGET, backup)
    print(f'  Бэкап: {os.path.basename(backup)}')
    return backup


def is_caption(text):
    """Подпись рисунка / таблицы — оставляем тире как есть."""
    t = text.strip()
    return (t.startswith('Рисунок ') or t.startswith('Таблица ')
            or t.startswith('Продолжение таблицы'))


def is_in_sources_zone(idx, sources_start, sources_end):
    return sources_start is not None and sources_start <= idx <= sources_end


def replace_text_keep_style(p, new_text):
    """Заменяет текст параграфа, сохраняя стиль первого run."""
    if p.text == new_text:
        return False
    runs = p.runs
    for r in runs[1:]:
        r._r.getparent().remove(r._r)
    if runs:
        runs[0].text = new_text
    else:
        p.add_run(new_text)
    return True


def fix_paragraph_text(text):
    """Возвращает (новый_текст, [список_правок])."""
    fixes = []
    new = text

    # 1) "(рисунок N)" → ", как показано на рисунке N"
    def _rep_fig(m):
        n = m.group(1)
        fixes.append(f'(рисунок {n})')
        return f', как показано на рисунке {n}'
    new = re.sub(r'\s*\(\s*рисунок\s+(\d+)\s*\)', _rep_fig, new)
    # Тот же случай но с заглавной (на всякий)
    new = re.sub(r'\s*\(\s*Рисунок\s+(\d+)\s*\)',
                  lambda m: f', как показано на рисунке {m.group(1)}',
                  new)

    # 2) Унификация тире: ' - ' и ' – ' → ' — '
    before = new
    new = re.sub(r' - ', ' — ', new)
    new = re.sub(r' – ', ' — ', new)
    if new != before:
        fixes.append('тире')

    return new, fixes


def main():
    print('=== Безопасные текстовые правки ===\n')
    make_backup()
    doc = Document(TARGET)

    # Найдём границы списка источников (там тире НЕ трогаем —
    # стандартный формат «— М.: Издательство, 2023.»)
    paragraphs = list(doc.paragraphs)
    sources_start = None
    sources_end = None
    for i, p in enumerate(paragraphs):
        t = p.text.strip()
        if t == 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ':
            sources_start = i
        if sources_start and i > sources_start and t.startswith(('ПРИЛОЖЕНИЕ',
                                                                  'ПЕРЕЧЕНЬ')):
            sources_end = i - 1
            break
    if sources_end is None and sources_start:
        sources_end = len(paragraphs) - 1

    counts = {'(рисунок N)': 0, 'тире': 0, 'caption-skipped': 0,
              'sources-skipped': 0}

    for i, p in enumerate(paragraphs):
        if is_caption(p.text):
            counts['caption-skipped'] += 1
            continue
        if is_in_sources_zone(i, sources_start, sources_end):
            counts['sources-skipped'] += 1
            continue

        new, fixes = fix_paragraph_text(p.text)
        if new != p.text:
            replace_text_keep_style(p, new)
            for f in fixes:
                if f.startswith('(рисунок'):
                    counts['(рисунок N)'] += 1
                else:
                    counts[f] += 1

    # Также — в таблицах, но там тире обычно правильные; правим только
    # «(рисунок N)»
    for tab in doc.tables:
        for row in tab.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    new = re.sub(r'\s*\(\s*рисунок\s+(\d+)\s*\)',
                                  lambda m: f', как показано на рисунке {m.group(1)}',
                                  p.text)
                    if new != p.text:
                        replace_text_keep_style(p, new)
                        counts['(рисунок N)'] += 1

    doc.save(TARGET)
    print('--- Результаты ---')
    for k, v in counts.items():
        print(f'  {k}: {v}')
    print(f'\n✓ Сохранено: {os.path.basename(TARGET)}')


if __name__ == '__main__':
    main()
