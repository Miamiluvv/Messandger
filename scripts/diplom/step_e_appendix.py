# -*- coding: utf-8 -*-
"""Перенос рисунков 6 (ER) и 20 (алгоритм чата) в приложение А."""
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .helpers import (set_paragraph_text, insert_paragraphs_after,
                      find_paragraph_starts_with, find_template_paragraph,
                      insert_paragraph_after)


APPENDIX_NOTES = [
    "А.6 ER-диаграмма базы данных приведена на рисунке А.6. Диаграмма "
    "отражает все основные сущности системы и связи между ними: "
    "пользователи (User), подразделения (Department), чаты (Chat), "
    "участники чатов (ChatMember), сообщения (Message), вложения "
    "(Attachment), уведомления (Notification), запланированные "
    "сообщения, опросы (Poll) и варианты ответов (PollOption), "
    "журнал аудита (AuditLog).",
]

APPENDIX_FIGURES_1 = [
    "",
    "Рисунок А.6 — ER-диаграмма базы данных",
    "",
]

APPENDIX_NOTES_2 = [
    "А.7 Полная блок-схема алгоритма создания группового чата "
    "приведена на рисунке А.7. Алгоритм охватывает все ветвления: "
    "проверку прав на создание чата заданного типа, валидацию состава "
    "участников, автоматическое повышение роли начальников управлений "
    "и их заместителей до администратора чата, обработку ошибок при "
    "добавлении участников и формирование результирующего "
    "идентификатора чата.",
]

APPENDIX_FIGURES_2 = [
    "",
    "Рисунок А.7 — Блок-схема алгоритма создания группового чата",
]


def run(doc):
    # Замены формулировок в основном тексте — рисунки оставляем на месте,
    # но добавляем перекрёстную ссылку на приложение.
    pairs = [
        ("Представленная на рисунке 6 ER-диаграмма",
         "ER-диаграмма приведена в Приложении А (рисунок А.6); она"),

        ("Алгоритм, представленный на рисунке 20,",
         "Алгоритм приведён в Приложении А (рисунок А.7); он"),

        ("(рисунок 6)", "(Приложение А, рисунок А.6)"),
        ("(рисунок 20)", "(Приложение А, рисунок А.7)"),
    ]
    changed = 0
    for p in doc.paragraphs:
        old = p.text
        new = old
        for s, r in pairs:
            new = new.replace(s, r)
        if new != old:
            set_paragraph_text(p, new)
            changed += 1

    # Вставить блок в Приложение А — пробуем несколько вариантов якоря
    anchor = find_paragraph_starts_with(doc,
        "Дополнительно в приложении представлен снимок экрана")
    if anchor is None:
        # В новой версии файла такого абзаца нет — берём последний параграф,
        # начинающийся с "Рисунок А."
        last_a = None
        for p in doc.paragraphs:
            if p.text.strip().startswith("Рисунок А."):
                last_a = p
        anchor = last_a
    if anchor is None:
        # Совсем крайний fallback — последний параграф документа
        anchor = doc.paragraphs[-1] if doc.paragraphs else None
    if anchor is not None:
        body_template = find_template_paragraph(
            doc, lambda p: p.text.startswith("Архитектура системы спроектирована"))
        caption_template = find_template_paragraph(
            doc, lambda p: p.text.strip().startswith("Рисунок 1 — "))

        last = insert_paragraphs_after(anchor, APPENDIX_NOTES,
                                         clone_from=body_template)
        for line in APPENDIX_FIGURES_1:
            last = insert_paragraph_after(last, line,
                                            clone_from=caption_template)
        last = insert_paragraphs_after(last, APPENDIX_NOTES_2,
                                         clone_from=body_template)
        for line in APPENDIX_FIGURES_2:
            last = insert_paragraph_after(last, line,
                                            clone_from=caption_template)
    print(f"[E] Рисунки 6 и 20 продублированы в Приложение А (изменено {changed} абз.)")
