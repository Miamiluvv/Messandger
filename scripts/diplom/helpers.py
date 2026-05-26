# -*- coding: utf-8 -*-
"""Утилиты работы с python-docx."""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement


def set_paragraph_text(p, new_text, *, bold=False, italic=False, size_pt=None):
    """Полная замена текста параграфа, формат первого rana сохраняется."""
    if p.runs:
        first = p.runs[0]
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
        first.text = new_text
        if bold:
            first.bold = True
        if italic:
            first.italic = True
        if size_pt:
            first.font.size = Pt(size_pt)
    else:
        run = p.add_run(new_text)
        run.bold = bold
        run.italic = italic
        if size_pt:
            run.font.size = Pt(size_pt)


def insert_paragraph_after(paragraph, text="", *, bold=False, italic=False,
                            alignment=None, indent=True, clone_from=None):
    """Вставить параграф после данного.

    По умолчанию форматирование наследуется от стиля Normal документа.
    Если задан clone_from — копируются все ключевые свойства параграфа
    (выравнивание, отступы, межстрочный, размер шрифта первого рана),
    что обеспечивает визуальную идентичность с соседними абзацами."""
    new_el = OxmlElement("w:p")
    paragraph._element.addnext(new_el)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_el, paragraph._parent)

    if clone_from is not None:
        src_pf = clone_from.paragraph_format
        dst_pf = p.paragraph_format
        # Копируем геометрию абзаца
        if src_pf.first_line_indent is not None:
            dst_pf.first_line_indent = src_pf.first_line_indent
        if src_pf.left_indent is not None:
            dst_pf.left_indent = src_pf.left_indent
        if src_pf.line_spacing is not None:
            dst_pf.line_spacing = src_pf.line_spacing
        if src_pf.space_before is not None:
            dst_pf.space_before = src_pf.space_before
        if src_pf.space_after is not None:
            dst_pf.space_after = src_pf.space_after
        if clone_from.alignment is not None:
            p.alignment = clone_from.alignment
        # Стиль
        try:
            p.style = clone_from.style
        except Exception:
            pass

    if text:
        run = p.add_run(text)
        if clone_from is not None and clone_from.runs:
            src_run = clone_from.runs[0]
            if src_run.font.name:
                run.font.name = src_run.font.name
            if src_run.font.size:
                run.font.size = src_run.font.size
        if bold:
            run.bold = True
        if italic:
            run.italic = True

    if alignment is not None:
        p.alignment = alignment
    return p


def insert_paragraphs_after(paragraph, texts, *, clone_from=None, **kwargs):
    """Вставить серию параграфов после данного. Возвращает последний."""
    last = paragraph
    for t in texts:
        last = insert_paragraph_after(last, t, clone_from=clone_from, **kwargs)
    return last


def find_template_paragraph(doc, predicate):
    """Найти первый параграф-образец, по которому потом клонируем
    форматирование. predicate(p) → bool."""
    for p in doc.paragraphs:
        if predicate(p):
            return p
    return None


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def find_paragraph(doc, predicate):
    """Первый параграф, удовлетворяющий predicate(text), либо None."""
    for p in doc.paragraphs:
        if predicate(p.text):
            return p
    return None


def find_paragraph_starts_with(doc, prefix):
    for p in doc.paragraphs:
        if p.text.startswith(prefix):
            return p
    return None


def find_index_of_paragraph(doc, paragraph):
    """Найти 0-индекс данного параграфа в doc.paragraphs по XML-элементу
    (стандартный .index() ломается из-за пересоздания обёрток)."""
    target_el = paragraph._element
    for i, p in enumerate(doc.paragraphs):
        if p._element is target_el:
            return i
    return -1


def previous_paragraph(doc, paragraph):
    idx = find_index_of_paragraph(doc, paragraph)
    if idx <= 0:
        return paragraph
    return doc.paragraphs[idx - 1]


def replace_in_table_cells(doc, mapping):
    """Текстовая замена в ячейках таблиц по словарю {old: new}."""
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    txt = p.text
                    new = txt
                    for old, repl in mapping.items():
                        new = new.replace(old, repl)
                    if new != txt:
                        set_paragraph_text(p, new)
