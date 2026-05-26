"""Утилита для дампа структуры .docx — параграфы, стили, рисунки, таблицы."""
import sys
import os
from docx import Document
from docx.oxml.ns import qn

def dump(path, out_path):
    doc = Document(path)
    lines = [f"=== {os.path.basename(path)} ==="]
    lines.append(f"Параграфов: {len(doc.paragraphs)}, Таблиц: {len(doc.tables)}")
    lines.append("")

    # Параграфы с номерами
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name if p.style else "?"
        # Проверка наличия картинки
        has_img = bool(p._element.findall('.//' + qn('w:drawing')))
        marker = " [IMG]" if has_img else ""
        if text or has_img:
            lines.append(f"[{i:04d}] ({style}){marker}: {text[:200]}")

    lines.append("")
    lines.append(f"=== Таблицы ({len(doc.tables)}) ===")
    for ti, t in enumerate(doc.tables):
        lines.append(f"-- Таблица {ti}: {len(t.rows)}x{len(t.columns) if t.rows else 0}")
        for ri, row in enumerate(t.rows[:3]):
            cells = [c.text.strip()[:50] for c in row.cells]
            lines.append(f"   [{ri}] {cells}")
        if len(t.rows) > 3:
            lines.append(f"   ... ещё {len(t.rows)-3} строк")

    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"OK: {out_path}")

if __name__ == "__main__":
    base = r"c:\Users\dalga\CascadeProjects\windsurf-project"
    files = [
        ("ДИПЛОМ_Махмудова_А_Д.docx", "diplom_dump.txt"),
        ("Махмудова ТЗ.docx", "tz_dump.txt"),
        ("курсовая_работа_Махмудова.docx", "kursovaya_dump.txt"),
        ("3 РАЗДЕЛ ОТДЕЛЬНО.docx", "razdel3_dump.txt"),
    ]
    out_dir = os.path.join(base, "scripts", "dumps")
    os.makedirs(out_dir, exist_ok=True)
    for src, dst in files:
        try:
            dump(os.path.join(base, src), os.path.join(out_dir, dst))
        except Exception as e:
            print(f"FAIL {src}: {e}")
