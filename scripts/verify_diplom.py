# -*- coding: utf-8 -*-
"""Сверка результата правки диплома."""
import os
from docx import Document

BASE = r"c:\Users\dalga\CascadeProjects\windsurf-project"
DST = os.path.join(BASE, "ДИПЛОМ_Махмудова_А_Д_v2.docx")

doc = Document(DST)
print(f"Параграфов в v2: {len(doc.paragraphs)}, таблиц: {len(doc.tables)}\n")

# Проверки
checks = {
    "Тире — заменены": ("–", "не должно быть en-dash"),
    "Точка интеграции вставлена": (
        "точка интеграции с единой корпоративной системой", "+"),
    "Раздел 2.7.1 безопасности": (
        "2.7.1 Обеспечение информационной безопасности", "+"),
    "WebRTC описание": ("Подсистема аудио- и видеозвонков построена на технологии WebRTC", "+"),
    "Редактор фото": ("встроенный редактор фотографий", "+"),
    "Сервер в Москве": ("физически расположенном на территории города Москвы", "+"),
    "HSTS": ("HSTS", "+"),
    "Бэкапы": ("Резервные копии базы данных создаются автоматически каждые шесть часов", "+"),
    "Раздел 3.6 (влияние)": (
        "3.6 Влияние процессов работы с информацией", "+"),
    "Новые экономические числа": ("8 814 098", "+"),
    "Новый PP": ("5,25 месяца", "+"),
    "Новый ROI": ("128,7", "+"),
    "Источники: нормативные": ("Нормативно-правовые акты", "+"),
    "Источники: стандарты": ("Стандарты", "+"),
    "Источники: электронные": ("Электронные ресурсы и техническая документация", "+"),
    "Ссылка [4] на 152-ФЗ": ("152-ФЗ [4]", "+"),
    "Ссылка [9] на ПП 1119": ("№ 1119 [9]", "+"),
    "Замечание к рис. 19": ("[ЗАМЕЧАНИЕ К РИСУНКУ 19", "+"),
    "Приложение А, рисунок А.1": ("Приложении А, рисунок А.1", "+"),
}

text_all = "\n".join(p.text for p in doc.paragraphs)
for cell in (c for t in doc.tables for r in t.rows for c in r.cells):
    text_all += "\n" + cell.text

for name, (needle, expected) in checks.items():
    found = needle in text_all
    if expected == "+":
        status = "OK" if found else "FAIL"
    else:
        status = "OK" if not found else "FAIL (нашёл лишнее)"
    print(f"  [{status}] {name}: '{needle[:60]}...'")

print("\n--- Первые 5 источников ---")
in_sources = False
shown = 0
for p in doc.paragraphs:
    if p.text.startswith("СПИСОК ИСПОЛЬЗОВАННЫХ"):
        in_sources = True
        continue
    if in_sources and p.text.strip():
        print(f"  {p.text[:120]}")
        shown += 1
        if shown >= 8:
            break
