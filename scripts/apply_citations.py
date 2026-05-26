# -*- coding: utf-8 -*-
"""Расстановка ссылок [N] после упоминаний законов и стандартов.

Карта (после реструктуризации источников):
  [1]  Конституция РФ
  [2]  Гражданский кодекс
  [3]  Трудовой кодекс (статья 91)
  [4]  Налоговый кодекс (глава 34 — страховые взносы)
  [5]  КоАП РФ (статья 13.11)
  [6]  ФЗ № 149-ФЗ
  [7]  ФЗ № 152-ФЗ
  [8]  ФЗ № 176-ФЗ
  [9]  ПП РФ № 1119
  [10] Приказ ФСТЭК № 21
  [11] ГОСТ Р 7.0.97
  [12] ГОСТ 7.32-2017

Идемпотентно: если после маркера уже стоит [N], не дублируем.
"""
import os, re, shutil
from datetime import datetime
from docx import Document

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TARGET = os.path.join(ROOT, 'ДИПЛОМ_Махмудова_А_Д [HIGirw] [Mbuh0y].docx')


# (regex для маркера, номер источника, описание для отчёта)
# Регексы устроены так: после захвата маркера НЕ должно сразу идти \s*\[\d+\]
PATTERNS = [
    (r'(Федеральн\w+\s+закон\w*(?:\s+(?:от\s+\d[\d.]*))?\s+№\s*152-ФЗ)',
     7, '152-ФЗ'),
    (r'(Федеральн\w+\s+закон\w*(?:\s+(?:от\s+\d[\d.]*))?\s+№\s*149-ФЗ)',
     6, '149-ФЗ'),
    (r'(Федеральн\w+\s+закон\w*(?:\s+(?:от\s+\d[\d.]*))?\s+№\s*176-ФЗ)',
     8, '176-ФЗ'),
    (r'(Гражданск\w+\s+кодекс\w*\s+(?:Российской\s+Федерации))',
     2, 'ГК'),
    (r'(Налогов\w+\s+кодекс\w*(?:\s+Российской\s+Федерации)?)',
     4, 'НК'),
    (r'(Трудов\w+\s+кодекс\w*(?:\s+Российской\s+Федерации)?)',
     3, 'ТК'),
    (r'(Кодекс\w*\s+Российской\s+Федерации\s+об\s+административн\w+\s+правонарушени\w+)',
     5, 'КоАП'),
    (r'(статье\s+13\.11\s+(?:Кодекса)?)',
     5, 'КоАП ст.13.11'),
    (r'(Постановлени\w+\s+Правительства\s+(?:Российской\s+Федерации|РФ)?\s*(?:от\s+\d[\d.]*\s+)?№\s*1119)',
     9, 'ПП 1119'),
    (r'(Приказ\w*\s+ФСТЭК[^.]{0,40}№\s*21)',
     10, 'Приказ ФСТЭК 21'),
    (r'(ГОСТ\s+Р\s+7\.0\.97[\-–]2016)',
     11, 'ГОСТ 7.0.97'),
    (r'(ГОСТ\s+7\.32[\-–]2017)',
     12, 'ГОСТ 7.32'),
    (r'(Конституци\w+\s+Российской\s+Федерации)',
     1, 'Конституция'),
]


def make_backup():
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    backup = os.path.join(ROOT, f'_BACKUP_до_цитирований_{ts}.docx')
    shutil.copy2(TARGET, backup)
    print(f'  Бэкап: {os.path.basename(backup)}')


def replace_text_keep_style(p, new_text):
    if p.text == new_text:
        return False
    runs = p.runs
    for r in runs[1:]:
        r._r.getparent().remove(r._r)
    if runs:
        runs[0].text = new_text
    else:
        p.add_run(new_text)
    return True


def is_skip_paragraph(text):
    t = text.strip()
    return (t.startswith('Рисунок ') or t.startswith('Таблица ')
            or t.startswith('Продолжение таблицы'))


def main():
    print('=== Расстановка ссылок [N] на источники ===\n')
    make_backup()
    doc = Document(TARGET)

    # Границы списка источников — там НЕ расставляем
    paragraphs = list(doc.paragraphs)
    src_start = None
    src_end = len(paragraphs)
    for i, p in enumerate(paragraphs):
        t = p.text.strip()
        if t == 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ':
            src_start = i
        if src_start is not None and i > src_start and \
                t.startswith(('ПРИЛОЖЕНИЕ', 'ПЕРЕЧЕНЬ')):
            src_end = i
            break

    counts = {desc: 0 for _, _, desc in PATTERNS}

    for i, p in enumerate(paragraphs):
        if src_start is not None and src_start <= i < src_end:
            continue
        if is_skip_paragraph(p.text):
            continue
        text = p.text
        if not text.strip():
            continue
        original = text
        for regex, num, desc in PATTERNS:
            # Идемпотентность: после маркера не должно быть [N]
            # Делаем функцию для замены
            ref = f' [{num}]'
            def _replace(m, ref=ref):
                end = m.end()
                # Если уже есть [N] сразу после — не добавляем
                tail = text[end:end + 8]
                if re.match(r'\s*\[\d+\]', tail):
                    return m.group(0)
                counts[desc] += 1
                return m.group(0) + ref

            new_text = re.sub(regex, _replace, text)
            text = new_text

        if text != original:
            replace_text_keep_style(p, text)

    doc.save(TARGET)
    print('\n--- Расстановка ---')
    total = 0
    for desc, n in counts.items():
        if n > 0:
            print(f'  [{desc}] добавлено: {n}')
            total += n
    print(f'\n✓ Всего вставлено ссылок: {total}')


if __name__ == '__main__':
    main()
